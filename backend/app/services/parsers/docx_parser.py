"""DOCX Parser — 使用 python-docx 解析 Word 文档"""

import uuid
import hashlib
import io
import re
from pathlib import Path
from typing import Optional
from xml.etree import ElementTree as ET

from loguru import logger

from app.services.parsers.base import (
    BaseParser, ParserRegistry, UDRAsset, UDRBlock, UnifiedDocument,
)

try:
    import docx
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    logger.warning("python-docx not installed; DOCX parsing will be limited")

# Word 命名空间
WORD_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}

# Heading 样式到层级的映射
HEADING_LEVEL_MAP = {
    "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
    "Heading 4": 4, "Heading 5": 5, "Heading 6": 6,
    "heading 1": 1, "heading 2": 2, "heading 3": 3,
    "heading 4": 4, "heading 5": 5, "heading 6": 6,
    "Titre 1": 1, "Titre 2": 2, "Titre 3": 3,
}


def _get_heading_level(style_name: Optional[str]) -> Optional[int]:
    """从样式名称推断标题层级"""
    if not style_name:
        return None
    # 精确匹配
    if style_name in HEADING_LEVEL_MAP:
        return HEADING_LEVEL_MAP[style_name]
    # 正则匹配 e.g. "Heading1", "heading1"
    m = re.match(r"[Hh]eading\s*(\d)", style_name)
    if m:
        return int(m.group(1))
    return None


def _run_to_markdown(run) -> str:
    """将单个 run 转为 markdown 文本，保留加粗/斜体等样式"""
    text = run.text or ""
    if not text:
        return ""
    if run.bold and run.italic:
        text = f"***{text}***"
    elif run.bold:
        text = f"**{text}**"
    elif run.italic:
        text = f"*{text}*"
    if run.underline:
        text = f"<u>{text}</u>"
    if run.strike:
        text = f"~~{text}~~"
    return text


def _extract_hyperlinks(paragraph) -> list[dict]:
    """从段落 XML 中提取超链接"""
    links = []
    try:
        p_elem = paragraph._element
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                 "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
        for hl in p_elem.findall(".//w:hyperlink", nsmap):
            rid = hl.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            link_text = "".join(node.text or "" for node in hl.findall(".//w:t", nsmap))
            url = ""
            if rid:
                try:
                    rel = paragraph.part.rels[rid]
                    url = rel.target_ref
                except Exception:
                    pass
            links.append({"text": link_text, "url": url})
    except Exception as e:
        logger.debug(f"Failed to extract hyperlinks: {e}")
    return links


def _paragraph_to_text(paragraph) -> str:
    """将段落转为 markdown 文本"""
    parts = []
    for run in paragraph.runs:
        parts.append(_run_to_markdown(run))
    return "".join(parts)


def _table_to_markdown_and_data(table) -> tuple[str, dict]:
    """将表格转为 Markdown 文本和 structured_data"""
    rows_data = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cells.append(cell.text.strip().replace("\n", " "))
        rows_data.append(cells)

    if not rows_data:
        return "", {"headers": [], "rows": []}

    headers = rows_data[0]
    data_rows = rows_data[1:] if len(rows_data) > 1 else []

    # Markdown 表格
    lines = ["| " + " | ".join(headers) + " |"]
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in data_rows:
        # 补齐列数
        padded = row + [""] * (len(headers) - len(row))
        lines.append("| " + " | ".join(padded[:len(headers)]) + " |")
    table_text = "\n".join(lines)

    return table_text, {"headers": headers, "rows": data_rows}


def _extract_comments(doc) -> list[dict]:
    """提取文档批注"""
    comments = []
    try:
        # python-docx 不直接支持批注，需解析 XML
        if not doc.part.package.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
        ):
            return comments
    except Exception:
        pass

    try:
        comments_part = None
        for rel in doc.part.rels.values():
            if "comments" in rel.reltype:
                comments_part = rel.target_part
                break
        if comments_part is None:
            return comments

        root = ET.fromstring(comments_part.blob)
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for comment_elem in root.findall(".//w:comment", nsmap):
            author = comment_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author", "")
            date = comment_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date", "")
            comment_id = comment_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", "")
            texts = []
            for t in comment_elem.findall(".//w:t", nsmap):
                if t.text:
                    texts.append(t.text)
            comments.append({
                "id": comment_id,
                "author": author,
                "date": date,
                "text": "".join(texts),
            })
    except Exception as e:
        logger.debug(f"Failed to extract comments: {e}")
    return comments


def _extract_footnotes(doc) -> list[dict]:
    """提取脚注"""
    footnotes = []
    try:
        footnotes_part = None
        for rel in doc.part.rels.values():
            if "footnotes" in rel.reltype:
                footnotes_part = rel.target_part
                break
        if footnotes_part is None:
            return footnotes

        root = ET.fromstring(footnotes_part.blob)
        nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for fn_elem in root.findall(".//w:footnote", nsmap):
            fn_id = fn_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", "")
            fn_type = fn_elem.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type", "")
            # 跳过分隔符脚注
            if fn_type in ("separator", "continuationSeparator"):
                continue
            texts = []
            for t in fn_elem.findall(".//w:t", nsmap):
                if t.text:
                    texts.append(t.text)
            footnotes.append({
                "id": fn_id,
                "text": "".join(texts),
            })
    except Exception as e:
        logger.debug(f"Failed to extract footnotes: {e}")
    return footnotes


def _extract_image_from_run(run, doc_id: str, img_idx: int) -> Optional[UDRAsset]:
    """尝试从 run 中提取图片并上传到 MinIO"""
    try:
        inline = run._element.findall(
            ".//wp:inline", WORD_NS
        )
        if not inline:
            return None

        for drawing in run._element.findall(".//w:drawing", WORD_NS):
            for blip in drawing.findall(".//a:blip", WORD_NS):
                embed_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                if not embed_id:
                    continue
                try:
                    image_part = run.part.rels[embed_id].target_part
                    image_data = image_part.blob
                    content_type = image_part.content_type

                    # 尝试上传到 MinIO
                    asset_id = f"asset_{uuid.uuid4().hex[:8]}"
                    ext = content_type.split("/")[-1] if "/" in content_type else "png"
                    if ext == "jpeg":
                        ext = "jpg"
                    object_name = f"parsed_assets/{doc_id}/{asset_id}.{ext}"

                    try:
                        from app.services.storage import minio_storage
                        minio_storage.upload_file(
                            object_name=object_name,
                            data=image_data,
                            content_type=content_type,
                        )
                        asset_uri = f"minio://{minio_storage.client._base_url.host}/{object_name}"
                    except Exception as e:
                        logger.warning(f"Failed to upload image to MinIO, using inline: {e}")
                        asset_uri = f"inline://{asset_id}"

                    checksum = hashlib.md5(image_data).hexdigest()
                    return UDRAsset(
                        asset_id=asset_id,
                        asset_type="image",
                        asset_uri=asset_uri,
                        mime_type=content_type,
                        file_size=len(image_data),
                        checksum=checksum,
                        metadata={"source": "docx_embedded"},
                    )
                except Exception as e:
                    logger.debug(f"Failed to extract image from run: {e}")
                    continue
    except Exception as e:
        logger.debug(f"Image extraction error: {e}")
    return None


class DOCXParser(BaseParser):
    supported_mime_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]
    supported_extensions = [".docx"]

    def parse(self, file_path: str, options: Optional[dict] = None) -> UnifiedDocument:
        if not HAS_PYTHON_DOCX:
            return self._fallback_parse(file_path)

        doc_id = f"doc_{uuid.uuid4().hex[:12]}"
        path = Path(file_path)
        filename = path.name
        options = options or {}

        blocks: list[UDRBlock] = []
        assets: list[UDRAsset] = []
        all_hyperlinks: list[dict] = []

        try:
            doc = docx.Document(file_path)
            core_props = doc.core_properties

            # 提取批注
            comments = _extract_comments(doc)
            # 提取脚注
            footnotes = _extract_footnotes(doc)

            # 跟踪 section_path
            heading_stack: list[tuple[int, str]] = []
            section_path = ""

            img_idx = 0
            for para_idx, para in enumerate(doc.paragraphs):
                style_name = para.style.name if para.style else ""
                heading_level = _get_heading_level(style_name)
                para_text = _paragraph_to_text(para)

                # 提取超链接
                links = _extract_hyperlinks(para)
                if links:
                    all_hyperlinks.extend(links)

                # 提取图片
                for run in para.runs:
                    asset = _extract_image_from_run(run, doc_id, img_idx)
                    if asset:
                        img_idx += 1
                        assets.append(asset)
                        blocks.append(UDRBlock(
                            block_id=f"{doc_id}_b{len(blocks):04d}",
                            type="image",
                            text=f"[Image: {asset.asset_id}]",
                            section_path=section_path,
                            asset_uri=asset.asset_uri,
                            metadata={"asset_id": asset.asset_id, "mime_type": asset.mime_type},
                        ))

                # 空段落跳过
                if not para_text.strip():
                    continue

                # 标题
                if heading_level is not None:
                    while heading_stack and heading_stack[-1][0] >= heading_level:
                        heading_stack.pop()
                    heading_stack.append((heading_level, para_text.strip()))
                    section_path = " > ".join(h[1] for h in heading_stack)
                    blocks.append(UDRBlock(
                        block_id=f"{doc_id}_b{len(blocks):04d}",
                        type="heading",
                        text=para_text.strip(),
                        level=heading_level,
                        section_path=section_path,
                    ))
                    continue

                # 正文段落
                blocks.append(UDRBlock(
                    block_id=f"{doc_id}_b{len(blocks):04d}",
                    type="paragraph",
                    text=para_text.strip(),
                    section_path=section_path,
                ))

            # 提取表格
            for t_idx, table in enumerate(doc.tables):
                try:
                    table_text, structured_data = _table_to_markdown_and_data(table)
                    blocks.append(UDRBlock(
                        block_id=f"{doc_id}_b{len(blocks):04d}",
                        type="table",
                        text=table_text,
                        section_path=section_path,
                        structured_data=structured_data,
                    ))
                except Exception as e:
                    logger.warning(f"Failed to extract table {t_idx}: {e}")

            # 批注作为 annotation blocks
            for comment in comments:
                blocks.append(UDRBlock(
                    block_id=f"{doc_id}_b{len(blocks):04d}",
                    type="annotation",
                    text=comment.get("text", ""),
                    section_path=section_path,
                    metadata={
                        "comment_id": comment.get("id", ""),
                        "author": comment.get("author", ""),
                        "date": comment.get("date", ""),
                    },
                ))

            # 脚注作为 metadata blocks
            for fn in footnotes:
                blocks.append(UDRBlock(
                    block_id=f"{doc_id}_b{len(blocks):04d}",
                    type="metadata",
                    text=fn.get("text", ""),
                    section_path=section_path,
                    metadata={"footnote_id": fn.get("id", ""), "footnote": True},
                ))

            # 超链接汇总
            link_metadata = {}
            if all_hyperlinks:
                link_metadata["hyperlinks"] = all_hyperlinks

            return UnifiedDocument(
                document_id=doc_id,
                source={
                    "filename": filename,
                    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "source_uri": str(path.absolute()),
                },
                metadata={
                    "title": core_props.title or filename,
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "keywords": core_props.keywords or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                    "comment_count": len(comments),
                    "footnote_count": len(footnotes),
                    **link_metadata,
                },
                blocks=blocks,
                assets=assets,
            )

        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            return self._fallback_parse(file_path)

    @staticmethod
    def _fallback_parse(file_path: str) -> UnifiedDocument:
        """python-docx 不可用或解析失败时的兜底方案"""
        doc_id = f"doc_{uuid.uuid4().hex[:12]}"
        path = Path(file_path)
        return UnifiedDocument(
            document_id=doc_id,
            source={
                "filename": path.name,
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "source_uri": str(path.absolute()),
            },
            metadata={"title": path.name, "warning": "DOCX parsing library not available"},
            blocks=[
                UDRBlock(
                    block_id=f"{doc_id}_b0000",
                    type="paragraph",
                    text=f"[DOCX file: {path.name} — install python-docx for full parsing]",
                ),
            ],
        )


# 注册
ParserRegistry.register(DOCXParser())
